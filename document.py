# -*- coding: utf-8 -*-
import seaborn as sns
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from statsmodels.tsa.stattools import adfuller
from docx.oxml.ns import qn


csv_path='result.csv'
doc_path='分析报告.doc'
roll_path='roll.png'
samplesize=20000
pair_path='pair.png'
ma_path='moving_log_avg.png'

#载入dataframe
df = pd.read_csv(csv_path, header=0, sep=',', names=['time', 'score'])
lens = len(df) + 1
df['time'] = np.arange(1, lens)
df['score_cumsum'] = df['score'].cumsum()
df.apply(pd.to_numeric, errors='ignore')
dsc = df['score_cumsum']
df['dsc_log'] = np.log(dsc)

#采样画pair图
def savePair(df,samplesize=20000):
    df1 = df.sample(samplesize)
    sns.set(style="ticks")
    sns.set_context("paper")
    sns.pairplot(df1)
    plt.title('Pair Graph')
    plt.savefig(pair_path)

#画滑动平均图，默认12阶
def saveMovg(df,step=12):
    moving_avg = pd.rolling_mean(df['dsc_log'], step)
    sns.jointplot(x=df['time'],y=moving_avg,color='red')
    plt.title('Moving Average')
    plt.savefig(ma_path)

# 创建word文档
def createDoc(pairfig,rollfig,mafig):
    document=Document()
    p=document.add_heading(u'训练情况分析报告',0)
    p=document.add_heading(u'一、训练分数描述性统计',level=1)
    document.add_picture(pairfig,width=Inches(6.0))
    p=document.add_heading(u'二、训练分数12阶滑动平均数图',level=1)
    document.add_picture(mafig,width=Inches(6.0))
    p=document.add_heading(u'三、训练分数均值、标准差移动平均线',level=1)
    document.add_picture(rollfig,width=Inches(6.0))
    document.save(doc_path)

testStationarity(df['dsc_log'])
savePair(df)
saveMovg(df)
createDoc(pair_path,roll_path,ma_path)